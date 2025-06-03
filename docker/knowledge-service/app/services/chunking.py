"""
Text chunking service for document processing
"""

import logging
from typing import List, Dict, Any
import re
from io import BytesIO
import uuid
import io

# Document processing imports
try:
    import PyPDF2
    from docx import Document
except ImportError:
    # These will be installed via requirements.txt
    pass

# Enhanced PDF processing imports
import fitz  # PyMuPDF
import pdfplumber

logger = logging.getLogger(__name__)


class ChunkingService:
    """Service for extracting text from documents and chunking"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        Initialize the chunking service
        
        Args:
            chunk_size: Maximum characters per chunk
            chunk_overlap: Overlap between consecutive chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        logger.info(f"📋 ChunkingService initialized - chunk_size: {chunk_size}, overlap: {chunk_overlap}")

    async def extract_text_from_file(self, file_content: bytes, content_type: str, filename: str) -> str:
        """
        Extract text from various file formats
        
        Args:
            file_content: Raw file bytes
            content_type: MIME type of the file
            filename: Original filename for format detection
            
        Returns:
            Extracted text content
        """
        try:
            # Determine file type
            filename_lower = filename.lower()
            
            if filename_lower.endswith('.pdf') or 'pdf' in content_type.lower():
                return await self._extract_from_pdf_enhanced(file_content, filename)
            elif filename_lower.endswith('.docx') or 'wordprocessingml' in content_type.lower():
                return await self._extract_from_docx(file_content)
            elif filename_lower.endswith(('.txt', '.text')) or 'text' in content_type.lower():
                return await self._extract_from_text(file_content)
            else:
                # Try to detect based on content
                if file_content.startswith(b'%PDF'):
                    return await self._extract_from_pdf_enhanced(file_content, filename)
                elif file_content.startswith(b'PK\x03\x04'):
                    return await self._extract_from_docx(file_content)
                else:
                    # Fallback to text
                    return await self._extract_from_text(file_content)
                    
        except Exception as e:
            logger.error(f"❌ Text extraction failed for {filename}: {e}")
            raise

    async def _extract_from_pdf_enhanced(self, file_content: bytes, filename: str) -> str:
        """
        Enhanced PDF text extraction using multiple methods for best results
        
        Uses PyMuPDF as primary, pdfplumber for tables, PyPDF2 as fallback
        """
        text_content = []
        
        try:
            # Method 1: PyMuPDF (fitz) - Best for most PDFs
            logger.info(f"🔍 Trying PyMuPDF extraction for {filename}")
            doc = fitz.open(stream=file_content, filetype="pdf")
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Extract text with layout preservation
                page_text = page.get_text("text")
                
                # Also try to extract tables if any
                tables = page.find_tables()
                table_text = ""
                for table in tables:
                    try:
                        table_data = table.extract()
                        # Convert table to readable text format
                        for row in table_data:
                            if row:  # Skip empty rows
                                table_text += " | ".join(str(cell) if cell else "" for cell in row) + "\n"
                        table_text += "\n"
                    except:
                        continue
                
                # Combine regular text and table text
                combined_text = page_text
                if table_text.strip():
                    combined_text += f"\n\n--- Tables on Page {page_num + 1} ---\n{table_text}"
                
                if combined_text.strip():
                    text_content.append(f"--- Page {page_num + 1} ---\n{combined_text}")
            
            doc.close()
            
            if text_content:
                full_text = "\n\n".join(text_content)
                logger.info(f"✅ PyMuPDF extracted {len(full_text)} characters from {filename}")
                return full_text
                
        except Exception as e:
            logger.warning(f"⚠️ PyMuPDF extraction failed for {filename}: {e}, trying pdfplumber")
        
        try:
            # Method 2: pdfplumber - Excellent for tables and structured content
            logger.info(f"🔍 Trying pdfplumber extraction for {filename}")
            
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_text = ""
                    
                    # Extract regular text
                    text = page.extract_text()
                    if text:
                        page_text += text
                    
                    # Extract tables with better formatting
                    tables = page.extract_tables()
                    if tables:
                        page_text += f"\n\n--- Tables on Page {page_num + 1} ---\n"
                        for table in tables:
                            for row in table:
                                if row:
                                    page_text += " | ".join(str(cell) if cell else "" for cell in row) + "\n"
                            page_text += "\n"
                    
                    if page_text.strip():
                        text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
            
            if text_content:
                full_text = "\n\n".join(text_content)
                logger.info(f"✅ pdfplumber extracted {len(full_text)} characters from {filename}")
                return full_text
                
        except Exception as e:
            logger.warning(f"⚠️ pdfplumber extraction failed for {filename}: {e}, falling back to PyPDF2")
        
        try:
            # Method 3: PyPDF2 - Fallback method
            logger.info(f"🔍 Trying PyPDF2 extraction for {filename}")
            return await self._extract_from_pdf_fallback(file_content)
            
        except Exception as e:
            logger.error(f"❌ All PDF extraction methods failed for {filename}: {e}")
            raise

    async def _extract_from_pdf_fallback(self, file_content: bytes) -> str:
        """Fallback PDF extraction using PyPDF2"""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text_content = []
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                except Exception as e:
                    text_content.append(f"--- Page {page_num + 1} ---\n[Error extracting text: {str(e)}]")
            
            full_text = "\n\n".join(text_content)
            logger.info(f"✅ PyPDF2 fallback extracted {len(full_text)} characters")
            return full_text
        
        except Exception as e:
            logger.error(f"❌ PyPDF2 fallback extraction failed: {e}")
            raise
    
    async def _extract_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            # Create a Document from bytes
            docx_file = BytesIO(file_content)
            doc = Document(docx_file)
            
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            full_text = "\n\n".join(text_content)
            logger.info(f"✅ Extracted {len(full_text)} characters from DOCX")
            return full_text
            
        except Exception as e:
            logger.error(f"❌ DOCX extraction failed: {e}")
            raise
    
    async def _extract_from_text(self, file_content: bytes) -> str:
        """Extract text from plain text file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'ascii']
            
            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    logger.info(f"✅ Decoded text file using {encoding}")
                    return text
                except UnicodeDecodeError:
                    continue
            
            # If all fail, use utf-8 with error handling
            text = file_content.decode('utf-8', errors='replace')
            logger.warning("Used UTF-8 with error replacement")
            return text
            
        except Exception as e:
            logger.error(f"❌ Text extraction failed: {e}")
            raise
    
    async def chunk_text(self, text: str, metadata: Dict[str, Any] = None) -> List[Dict[str, Any]]:
        """
        Split text into overlapping chunks
        
        Args:
            text: Input text to chunk
            metadata: Additional metadata to include with each chunk
            
        Returns:
            List of chunk dictionaries
        """
        if not text.strip():
            return []
        
        try:
            # Clean and normalize text
            text = self._clean_text(text)
            
            # Split into sentences for better chunk boundaries
            sentences = self._split_into_sentences(text)
            
            chunks = []
            current_chunk = ""
            current_length = 0
            
            for sentence in sentences:
                sentence_length = len(sentence)
                
                # If adding this sentence would exceed chunk size
                if current_length + sentence_length > self.chunk_size and current_chunk:
                    # Save current chunk
                    chunk_data = {
                        "content": current_chunk.strip(),
                        "metadata": metadata or {},
                        "length": len(current_chunk.strip())
                    }
                    chunks.append(chunk_data)
                    
                    # Start new chunk with overlap
                    if self.chunk_overlap > 0:
                        overlap_text = self._get_overlap_text(current_chunk, self.chunk_overlap)
                        current_chunk = overlap_text + " " + sentence
                        current_length = len(current_chunk)
                    else:
                        current_chunk = sentence
                        current_length = sentence_length
                else:
                    # Add sentence to current chunk
                    if current_chunk:
                        current_chunk += " " + sentence
                    else:
                        current_chunk = sentence
                    current_length += sentence_length
            
            # Add final chunk if it has content
            if current_chunk.strip():
                chunk_data = {
                    "content": current_chunk.strip(),
                    "metadata": metadata or {},
                    "length": len(current_chunk.strip())
                }
                chunks.append(chunk_data)
            
            logger.info(f"✅ Created {len(chunks)} chunks from {len(text)} characters")
            return chunks
            
        except Exception as e:
            logger.error(f"❌ Text chunking failed: {e}")
            raise
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove excessive newlines but preserve paragraph breaks
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences for better chunk boundaries"""
        # Simple sentence splitting - could be improved with proper NLP
        sentence_endings = r'[.!?]+(?:\s|$)'
        sentences = re.split(sentence_endings, text)
        
        # Filter out empty sentences
        sentences = [s.strip() for s in sentences if s.strip()]
        
        return sentences
    
    def _get_overlap_text(self, text: str, overlap_size: int) -> str:
        """Get the last portion of text for overlap"""
        if len(text) <= overlap_size:
            return text
        
        # Try to find a good break point (end of sentence)
        overlap_text = text[-overlap_size:]
        
        # Look for sentence boundaries in the overlap
        sentence_end = re.search(r'[.!?]+\s+', overlap_text)
        if sentence_end:
            return overlap_text[sentence_end.end():]
        
        # If no sentence boundary, look for word boundaries
        words = overlap_text.split()
        if len(words) > 1:
            return ' '.join(words[1:])  # Remove partial first word
        
        return overlap_text
    
    async def process_document(self, file_content: bytes, content_type: str, 
                             filename: str, metadata: Dict[str, Any] = None) -> Dict[str, Any]:
        """
        Complete document processing pipeline
        
        Args:
            file_content: Raw file bytes
            content_type: MIME type
            filename: Original filename
            metadata: Additional metadata
            
        Returns:
            Dictionary with document info and chunks
        """
        try:
            # Generate unique document ID
            document_id = str(uuid.uuid4())
            
            # Extract text
            text_content = await self.extract_text_from_file(file_content, content_type, filename)
            
            # Create base metadata
            doc_metadata = {
                "filename": filename,
                "content_type": content_type,
                "size": len(file_content),
                "text_length": len(text_content),
                **(metadata or {})
            }
            
            # Chunk the text
            chunks = await self.chunk_text(text_content, doc_metadata)
            
            # Add document ID to each chunk
            for chunk in chunks:
                chunk["document_id"] = document_id
                chunk["filename"] = filename
                chunk["content_type"] = content_type
            
            result = {
                "document_id": document_id,
                "filename": filename,
                "content_type": content_type,
                "text_content": text_content,
                "chunks": chunks,
                "metadata": doc_metadata
            }
            
            logger.info(f"✅ Processed document {filename}: {len(chunks)} chunks")
            return result
            
        except Exception as e:
            logger.error(f"❌ Document processing failed for {filename}: {e}")
            raise 